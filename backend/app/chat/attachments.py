"""Extracción de texto de documentos adjuntos al chat cognitivo.

El usuario adjunta un archivo en el Command Center; aquí lo convertimos a
texto plano para inyectarlo como contexto del turno actual del modelo. NO se
persiste el archivo: solo se extrae su texto y se devuelve al frontend, que lo
reenvía junto al mensaje.

Formatos soportados: PDF, Word (.docx), Excel (.xlsx/.xls), CSV/TSV y texto
plano (.txt/.md/.json/.xml/.log). Las imágenes NO se soportan: el modelo local
(gpt-oss-20b) es solo texto y no hay OCR instalado — se rechazan con un mensaje
honesto en vez de fingir que se leyeron.
"""

from __future__ import annotations

import csv
import io
import os

# Límite de tamaño del archivo subido (bytes). Evita OOM y subidas absurdas.
MAX_BYTES = 15 * 1024 * 1024  # 15 MB

# Tope de caracteres del texto extraído que se inyecta al modelo. El contexto
# de gpt-oss-20b es acotado (~24k tokens); truncamos para no reventarlo y para
# no ahogar la pregunta del usuario. ~1 token ≈ 4 chars → ~6k tokens de doc.
MAX_CHARS = 24_000

_TEXT_EXTS = {".txt", ".md", ".markdown", ".json", ".xml", ".log", ".yaml", ".yml", ".html", ".htm"}
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".heic"}


class AttachmentError(Exception):
    """Error legible (para el usuario) al procesar un adjunto."""


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    # último recurso: reemplazar bytes inválidos en vez de fallar
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    import pdfplumber  # import diferido: solo si llega un PDF

    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # También el texto de las tablas (fila por fila, celdas separadas por " | ")
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"# Hoja: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            if row is None:
                continue
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                out.append("\t".join(cells))
        out.append("")  # separador entre hojas
    wb.close()
    return "\n".join(out).strip()


def _extract_csv(data: bytes) -> str:
    text = _decode_text(data)
    # Normaliza el delimitador a tabulación para que quede compacto y legible.
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(io.StringIO(text), dialect)
    return "\n".join("\t".join(row) for row in reader).strip()


def extract(filename: str, data: bytes) -> dict:
    """Extrae el texto de un adjunto. Devuelve un dict serializable.

    Claves: name, kind, chars, truncated, text.
    Lanza AttachmentError con un mensaje legible si no se puede procesar.
    """
    name = (filename or "documento").strip()
    if not data:
        raise AttachmentError("El archivo está vacío.")
    if len(data) > MAX_BYTES:
        mb = len(data) / (1024 * 1024)
        raise AttachmentError(
            f"El archivo pesa {mb:.1f} MB; el máximo es {MAX_BYTES // (1024 * 1024)} MB."
        )

    ext = _ext(name)

    if ext in _IMAGE_EXTS:
        raise AttachmentError(
            "Las imágenes aún no se pueden leer: el modelo local es solo texto "
            "y no hay OCR. Adjunta un PDF, Word, Excel, CSV o texto."
        )

    try:
        if ext == ".pdf":
            kind, text = "pdf", _extract_pdf(data)
        elif ext == ".docx":
            kind, text = "word", _extract_docx(data)
        elif ext in (".xlsx", ".xlsm"):
            kind, text = "excel", _extract_xlsx(data)
        elif ext in (".csv", ".tsv"):
            kind, text = "csv", _extract_csv(data)
        elif ext in _TEXT_EXTS or ext == "":
            kind, text = "texto", _decode_text(data).strip()
        elif ext in (".xls", ".doc", ".ppt", ".pptx"):
            raise AttachmentError(
                f"El formato {ext} (Office antiguo) no se soporta directamente. "
                "Guárdalo como .docx/.xlsx o PDF y vuelve a adjuntarlo."
            )
        else:
            raise AttachmentError(
                f"No sé leer archivos {ext or 'sin extensión'}. Formatos soportados: "
                "PDF, Word (.docx), Excel (.xlsx), CSV y texto."
            )
    except AttachmentError:
        raise
    except Exception as exc:  # noqa: BLE001 — cualquier fallo del parser
        raise AttachmentError(f"No se pudo leer «{name}»: {exc}") from exc

    if not text:
        raise AttachmentError(
            f"«{name}» no contiene texto extraíble (¿es un escaneo/imagen?)."
        )

    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS].rstrip() + "\n\n[…documento truncado por longitud…]"

    return {
        "name": name,
        "kind": kind,
        "chars": len(text),
        "truncated": truncated,
        "text": text,
    }
