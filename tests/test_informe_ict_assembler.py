# tests/test_ict_report_assembler.py
import io

from docx import Document

from backend.app.aud.informe_cumplimiento_tributario import docx_assembler


def _ctx(**over):
    base = dict(
        firma_auditora="audit_consulting",
        razon_social="AXXISGASTRO CIA. LTDA.",
        ejercicio="2025",
        fecha_emision="27 de febrero de 2026",
        fecha_declaracion_ir="09 de abril de 2026",
        fecha_carga_sri="08 de julio de 2026",
        marco_contable="pymes",
        bloque_otros_asuntos="…no existen recomendaciones…",
        bloque_parte_iii="…no hemos identificado observaciones…",
    )
    base.update(over)
    return base


def _text(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    # doc.paragraphs no incluye texto dentro de tablas (gotcha de python-docx);
    # varios tokens (p.ej. fecha_declaracion_ir) viven en celdas de tabla.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_assemble_audit_consulting_rellena_tokens():
    out = docx_assembler.assemble(**_ctx())
    assert isinstance(out, bytes) and len(out) > 2000
    txt = _text(out)
    assert "AXXISGASTRO CIA. LTDA." in txt
    assert "27 de febrero de 2026" in txt
    assert "09 de abril de 2026" in txt
    assert "31 de diciembre de 2025" in txt
    assert "PYMES" in txt
    assert "{{" not in txt  # no quedan tokens sin rellenar


def test_assemble_partner_usa_su_plantilla():
    out = docx_assembler.assemble(**_ctx(firma_auditora="partner_auditing"))
    txt = _text(out)
    # dato horneado de Partner (socio)
    assert "Cristina Trujillo" in txt
    assert "{{" not in txt


def test_assemble_firma_invalida_lanza():
    import pytest
    with pytest.raises(ValueError):
        docx_assembler.assemble(**_ctx(firma_auditora="inexistente"))
