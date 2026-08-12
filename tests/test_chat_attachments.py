"""Tests de adjuntar documentos al chat cognitivo.

Cubren:
- Extractor: txt/csv se leen; imágenes se rechazan con mensaje honesto;
  formato desconocido se rechaza; truncado por longitud.
- Endpoint /chat/attachments/extract con .txt, .docx y .xlsx reales.
- Envío de mensaje con adjunto: el historial guarda el mensaje LIMPIO (con
  nota), pero el modelo recibe el texto del documento inyectado en ese turno.
"""

import io
import uuid

import pytest

from backend.app.auth import service as auth_service
from backend.app.auth.models import Role
from backend.app.chat import attachments
from backend.app.chat.providers import LLMResponse
from backend.app.db.session import SessionLocal, init_db


@pytest.fixture(autouse=True)
def _db():
    init_db()
    yield


def _mk(role: Role = Role.user):
    email = f"{role.value}-{uuid.uuid4().hex[:8]}@example.com"
    db = SessionLocal()
    try:
        auth_service.create_user(db, email=email, password="Sup3rSecret!", role=role)
    finally:
        db.close()
    return email, "Sup3rSecret!"


def _login(client, email, pw):
    r = client.post("/api/v1/auth/login", data={"username": email, "password": pw})
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


def _h(tok):
    return {"Authorization": f"Bearer {tok}"}


# ---- Extractor (unidad) ----------------------------------------------------

def test_extract_txt():
    out = attachments.extract("nota.txt", "Hola mundo\nSegunda línea".encode("utf-8"))
    assert out["kind"] == "texto"
    assert "Hola mundo" in out["text"]
    assert out["truncated"] is False


def test_extract_csv_normaliza():
    out = attachments.extract("datos.csv", b"a,b,c\n1,2,3\n")
    assert out["kind"] == "csv"
    assert "a\tb\tc" in out["text"]


def test_extract_image_rechazada():
    with pytest.raises(attachments.AttachmentError) as e:
        attachments.extract("foto.png", b"\x89PNG\r\n\x1a\n....")
    assert "imágenes" in str(e.value).lower()


def test_extract_formato_desconocido():
    with pytest.raises(attachments.AttachmentError):
        attachments.extract("cosa.zip", b"PK\x03\x04....")


def test_extract_trunca_documento_largo():
    grande = ("línea de relleno " * 5000).encode("utf-8")
    out = attachments.extract("largo.txt", grande)
    assert out["truncated"] is True
    assert out["chars"] <= attachments.MAX_CHARS + 60


def test_extract_vacio():
    with pytest.raises(attachments.AttachmentError):
        attachments.extract("vacio.txt", b"")


# ---- Endpoint /chat/attachments/extract ------------------------------------

def test_endpoint_extract_txt(client):
    email, pw = _mk()
    tok = _login(client, email, pw)
    r = client.post(
        "/api/v1/chat/attachments/extract",
        headers=_h(tok),
        files={"file": ("informe.txt", b"Ingresos 2026: 1.000.000 USD", "text/plain")},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["name"] == "informe.txt"
    assert "1.000.000" in body["text"]


def test_endpoint_extract_docx(client):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Acta de junta directiva")
    doc.add_paragraph("Acuerdo: aprobar el presupuesto.")
    buf = io.BytesIO()
    doc.save(buf)

    email, pw = _mk()
    tok = _login(client, email, pw)
    r = client.post(
        "/api/v1/chat/attachments/extract",
        headers=_h(tok),
        files={"file": ("acta.docx", buf.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    assert "Acta de junta" in r.json()["text"]


def test_endpoint_extract_xlsx(client):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Cuenta", "Saldo"])
    ws.append(["Caja", 1500])
    buf = io.BytesIO()
    wb.save(buf)

    email, pw = _mk()
    tok = _login(client, email, pw)
    r = client.post(
        "/api/v1/chat/attachments/extract",
        headers=_h(tok),
        files={"file": ("balance.xlsx", buf.getvalue(),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert r.status_code == 200, r.text
    text = r.json()["text"]
    assert "Cuenta" in text and "1500" in text


def test_endpoint_extract_imagen_422(client):
    email, pw = _mk()
    tok = _login(client, email, pw)
    r = client.post(
        "/api/v1/chat/attachments/extract",
        headers=_h(tok),
        files={"file": ("foto.jpg", b"\xff\xd8\xff\xe0junk", "image/jpeg")},
    )
    assert r.status_code == 422, r.text
    assert "imágenes" in r.json()["detail"].lower()


def test_endpoint_extract_requiere_auth(client):
    r = client.post(
        "/api/v1/chat/attachments/extract",
        files={"file": ("x.txt", b"hola", "text/plain")},
    )
    assert r.status_code in (401, 403)


# ---- Envío de mensaje con adjunto ------------------------------------------

def test_send_message_inyecta_adjunto_al_modelo_y_guarda_limpio(client, monkeypatch):
    captured: dict = {}

    def fake_complete(messages, system=None):
        captured["messages"] = messages
        return LLMResponse(content="Analizado.", model="test-model", tokens_in=1, tokens_out=1)

    from backend.app.chat import service as chat_service
    monkeypatch.setattr(chat_service, "chat_complete", fake_complete)

    email, pw = _mk()
    tok = _login(client, email, pw)
    conv = client.post("/api/v1/chat/conversations", headers=_h(tok), json={}).json()

    r = client.post(
        f"/api/v1/chat/conversations/{conv['id']}/messages",
        headers=_h(tok),
        json={
            "content": "Resume este documento.",
            "attachments": [
                {"name": "acta.txt", "text": "CONTENIDO_SECRETO_DEL_DOCUMENTO_XYZ"}
            ],
        },
    )
    assert r.status_code == 200, r.text

    # El modelo SÍ recibió el texto del documento inyectado.
    model_content = captured["messages"][-1]["content"]
    assert "CONTENIDO_SECRETO_DEL_DOCUMENTO_XYZ" in model_content
    assert "Resume este documento." in model_content

    # El historial guarda el mensaje LIMPIO (sin volcar el documento), con nota.
    detail = client.get(
        f"/api/v1/chat/conversations/{conv['id']}", headers=_h(tok)
    ).json()
    user_stored = detail["messages"][0]["content"]
    assert "CONTENIDO_SECRETO_DEL_DOCUMENTO_XYZ" not in user_stored
    assert "📎 Adjunto: acta.txt" in user_stored
