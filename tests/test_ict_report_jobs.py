# tests/test_ict_report_jobs.py
"""Test end-to-end de process_job (orquestador del Informe de Cumplimiento
Tributario) con PDFs reales.

Nota de adaptación (dos ajustes respecto al plan):

1. Extracción de texto del .docx consciente de tablas: `Document(...).paragraphs`
   NO recorre el texto dentro de celdas de tabla (gotcha de python-docx), y
   varios tokens (p.ej. las fechas) viven en celdas. Se usa `_docx_text()`,
   el mismo patrón ya usado en `tests/test_ict_report_assembler.py::_text()`.

2. Helper `_job()`: `ctx_service.create_client`/`create_project` son
   multi-tenant y requieren `organization_id` en este repo. Se sigue el
   patrón real de `tests/test_ict_report_service.py::_admin_and_project`
   (`ensure_user_has_organization` + `organization_id=user.organization_id`),
   devolviendo únicamente el `job.id` para evitar DetachedInstanceError.
"""

import json
import uuid
from pathlib import Path

import pytest
from docx import Document

from backend.app.auth import service as auth_service
from backend.app.auth.models import Role
from backend.app.aud.obligaciones_fiscales import file_storage
from backend.app.aud.informe_cumplimiento_tributario import jobs, service
from backend.app.context import service as ctx_service
from backend.app.db.session import SessionLocal, init_db

FIX = Path(__file__).parent / "fixtures" / "informe_cumplimiento_tributario"


@pytest.fixture(autouse=True)
def _db(tmp_path, monkeypatch):
    monkeypatch.setenv("AUD_OF_TMP_DIR", str(tmp_path))
    from importlib import reload
    from backend.app.core import config
    reload(config)
    reload(file_storage)
    init_db()
    yield


def _docx_text(path):
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts += [p.text for p in cell.paragraphs]
    return "\n".join(parts)


def _job():
    db = SessionLocal()
    try:
        tag = uuid.uuid4().hex[:6]
        user = auth_service.create_user(
            db, email=f"a-{tag}@ex.com", password="Sup3rSecret!", role=Role.admin
        )
        user = ctx_service.ensure_user_has_organization(db, user)
        client = ctx_service.create_client(
            db, organization_id=user.organization_id, name=f"C-{tag}"
        )
        project = ctx_service.create_project(
            db, organization_id=user.organization_id, client_id=client.id,
            name=f"P-{tag}", module_code="AUD",
        )
        job = service.create_job(
            db, user=user, project_id=project.id,
            cliente_name="AXXISGASTRO CIA. LTDA.", ejercicio="2025",
            firma_auditora="audit_consulting",
        )
        return job.id
    finally:
        db.close()


def test_process_job_genera_docx_con_datos_reales():
    jid = _job()
    jd = file_storage.create_job_dir(jid)
    file_storage.save_input(jd, "informe_auditoria_externa", "inf.pdf",
                            (FIX / "informe_auditoria_externa_axxis.pdf").read_bytes())
    file_storage.save_input(jd, "declaracion_ir", "f101.pdf",
                            (FIX / "f101_axxis.pdf").read_bytes())
    file_storage.save_input(jd, "params", "params.json", json.dumps({
        "fecha_carga_sri": "08 de julio de 2026",
        "hay_recomendaciones": False,
        "texto_recomendaciones": "",
        "override_fecha_emision": "",
        "override_marco_contable": "",
        "override_fecha_declaracion_ir": "",
    }).encode())

    jobs.process_job(jid)

    db = SessionLocal()
    try:
        from backend.app.aud.obligaciones_fiscales.models import ToolJob
        job = db.get(ToolJob, jid)
        assert job.status == "done", job.error_message
    finally:
        db.close()

    out = (file_storage.job_dir(jid) / "output.docx")
    assert out.exists()
    txt = _docx_text(out)
    assert "27 de febrero de 2026" in txt   # fecha emisión parseada
    assert "09 de abril de 2026" in txt     # fecha declaración parseada
    assert "31 de diciembre de 2025" in txt # cierre derivado
    assert "PYMES" in txt                    # marco detectado
    assert "{{" not in txt
